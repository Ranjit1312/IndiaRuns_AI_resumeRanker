"""gen_resumes.py — synthetic resume generator (Phase 2, Part C3).

Generates ~12 synthetic resumes aligned to a spread of the 10 gold JDs in
`data/eval_jds/*/` (reading each `jd_meta.yaml` must_haves/nice_to_haves and
`jd_profile.yaml` signals/role/locations to seed realistic content), across
three fit tiers per JD: **strong** / **borderline** / **weak**. These exist so
`redrob_ranker.fit.score_candidate` has a hand-authored spread of real-shaped
resumes to sanity-check against (strong should score meaningfully higher than
weak for its own JD).

Faithfulness / privacy:
  - Placeholder PII ONLY — name "XYZ Candidate", email "abc@example.com",
    phone "+91-XXXXXXXXXX". No real people are represented.
  - Deterministic: every resume's content is seeded off `random.seed(...)`
    derived from (jd_slug, tier), so re-running reproduces byte-identical
    output.
  - Exports plain text (always), DOCX (python-docx, a core dep), and PDF
    (fpdf2, optional — lazy-imported; if absent, PDF export is skipped with a
    warning and txt+docx are still written).

Output layout: `data/synthetic_resumes/<jd_slug>_<tier>/<layout>/resume.(txt|docx|pdf)`
plus a top-level `data/synthetic_resumes/INDEX.md`. `<layout>` is the
STRUCTURAL variety dimension added on top of content/tier — same content,
PII, and seed per (jd_slug, tier); only the rendering differs. See
`tools/resume_layouts.py` for the layout renderers and
`resume_layouts.LAYOUTS` for the full list (single_column, two_column,
table_heavy, ats_plain, image_only). `image_only` is a rasterized,
text-layer-free PDF built from the single_column PDF — it exists to exercise
the OCR path and has no txt/docx sibling.

Run:
    ./.venv/Scripts/python -m tools.gen_resumes
"""
from __future__ import annotations

import argparse
import os
import random
from dataclasses import dataclass, field

import yaml

from tools import resume_layouts

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
EVAL_JDS = os.path.join(ROOT, "data", "eval_jds")
OUT_DIR = os.path.join(ROOT, "data", "synthetic_resumes")

TIERS = ["strong", "borderline", "weak"]

# Four JDs spanning distinct domains (backend eng, security, applied-AI
# solutions architecture, business data science) x 3 tiers = 12 resumes —
# a representative slice of the 10-gold set rather than an exhaustive 30.
TARGET_SLUGS = [
    "stripe_backend-software-engineer",
    "amazon_security-engineer-aws",
    "nvidia_solutions-architect-ai-ml",
    "google-data-scientist",
]

PLACEHOLDER_NAME = "XYZ Candidate"
PLACEHOLDER_EMAIL = "abc@example.com"
PLACEHOLDER_PHONE = "+91-XXXXXXXXXX"

_COMPANY_POOL = ["Nimbus Systems", "Vertex Labs", "Cascade Works", "Brightloop Inc",
                 "Fernbridge Technologies", "Orbital Data Co", "Northstar Digital",
                 "Meridian Cloud", "Ironclad Analytics", "Quiet River Software"]
_UNRELATED_COMPANY_POOL = ["Sunrise Design Studio", "Metro Retail Group",
                           "Coastal Marketing Co", "Everyday Sales Partners"]
_TIER_1_SCHOOLS = ["IIT Delhi", "IIT Bombay", "BITS Pilani", "IIT Madras"]
_TIER_3_SCHOOLS = ["Regional Institute of Commerce", "State Polytechnic College"]


@dataclass
class JDContext:
    slug: str
    title: str
    company: str
    domain: str
    peak_years: float
    sigma_years: float
    preferred_locations: list
    acceptable_locations: list
    signals: list          # list of {id,label,query,evidence_regex}
    must_have_texts: list  # jd_meta must_haves[].text
    nice_have_texts: list
    in_domain_terms: list
    out_of_domain_terms: list


def load_jd_context(slug: str) -> JDContext:
    jd_dir = os.path.join(EVAL_JDS, slug)
    with open(os.path.join(jd_dir, "jd_profile.yaml"), encoding="utf-8") as f:
        prof = yaml.safe_load(f)
    meta_path = os.path.join(jd_dir, "jd_meta.yaml")
    meta = {}
    if os.path.exists(meta_path):
        with open(meta_path, encoding="utf-8") as f:
            meta = yaml.safe_load(f) or {}
    role = prof["role"]
    ie = role["ideal_experience"]
    return JDContext(
        slug=slug,
        title=role["title"],
        company=role.get("company", "") or "",
        domain=role.get("domain", "") or "",
        peak_years=float(ie["peak_years"]),
        sigma_years=float(ie["sigma_years"]),
        preferred_locations=list(prof["locations"].get("preferred") or []),
        acceptable_locations=list(prof["locations"].get("acceptable") or []),
        signals=prof.get("signals", []),
        must_have_texts=[m.get("text", "") for m in (meta.get("must_haves") or [])],
        nice_have_texts=[n.get("text", "") for n in (meta.get("nice_to_haves") or [])],
        in_domain_terms=list(prof.get("domain", {}).get("in_domain_terms") or []),
        out_of_domain_terms=list(prof.get("domain", {}).get("out_of_domain_terms") or []),
    )


# --------------------------------------------------------------------------- #
# content builders — one per tier, seeded off the JD's own signals
# --------------------------------------------------------------------------- #
def _signal_phrases(ctx: JDContext, n: int, rng: random.Random) -> list:
    labels = [s.get("label", s.get("id", "")) for s in ctx.signals]
    labels = [l.split("[")[0].strip() for l in labels if l]
    rng.shuffle(labels)
    return labels[:n]


def _bullet_from_signal(label: str, rng: random.Random) -> str:
    verbs = ["Led", "Owned", "Drove", "Delivered", "Built", "Designed and shipped"]
    return f"{rng.choice(verbs)} work directly demonstrating {label}, with measurable impact."


def _years_for_tier(ctx: JDContext, tier: str, rng: random.Random) -> float:
    if tier == "strong":
        return round(ctx.peak_years + rng.uniform(-0.5, 0.5), 1)
    if tier == "borderline":
        # near the edge of the sigma band -> partial fit
        offset = ctx.sigma_years * rng.choice([1.0, -1.0])
        return max(0.5, round(ctx.peak_years + offset, 1))
    # weak: far outside the ideal band, or drastically under-experienced
    return max(0.3, round(max(0.5, ctx.peak_years - ctx.sigma_years * 2.5), 1))


def _location_for_tier(ctx: JDContext, tier: str, rng: random.Random) -> str:
    if tier == "strong" and ctx.preferred_locations:
        return rng.choice(ctx.preferred_locations).title()
    if tier == "borderline" and ctx.acceptable_locations:
        return rng.choice(ctx.acceptable_locations).title()
    return "Nagpur"   # off-list city for every JD in this set -> location damp


def _career_history_strong(ctx: JDContext, yoe: float, rng: random.Random) -> list:
    n_roles = 2
    phrases = _signal_phrases(ctx, 6, rng)
    per_role = max(1, len(phrases) // n_roles)
    roles = []
    remaining_years = yoe
    for i in range(n_roles):
        dur_years = remaining_years / (n_roles - i)
        remaining_years -= dur_years
        bullets = phrases[i * per_role:(i + 1) * per_role] or phrases[:2]
        roles.append({
            "company": _COMPANY_POOL[(hash(ctx.slug) + i) % len(_COMPANY_POOL)],
            "title": ctx.title if i == 0 else f"{ctx.title} (Associate)",
            "years": max(0.5, round(dur_years, 1)),
            "bullets": [_bullet_from_signal(p, rng) for p in bullets] or
                       [f"Delivered core {ctx.domain or 'engineering'} work."],
            "is_current": i == 0,
        })
    return roles


def _career_history_borderline(ctx: JDContext, yoe: float, rng: random.Random) -> list:
    phrases = _signal_phrases(ctx, 3, rng)   # only ~half the signal surface
    return [
        {
            "company": _COMPANY_POOL[(hash(ctx.slug) + 1) % len(_COMPANY_POOL)],
            "title": f"{ctx.title.split(',')[0]} (Generalist)",
            "years": max(0.5, round(yoe * 0.6, 1)),
            "bullets": [_bullet_from_signal(p, rng) for p in phrases] or
                       ["Contributed to cross-functional delivery work."],
            "is_current": True,
        },
        {
            "company": _COMPANY_POOL[(hash(ctx.slug) + 2) % len(_COMPANY_POOL)],
            "title": "Generalist Associate",
            "years": max(0.5, round(yoe * 0.4, 1)),
            "bullets": ["Rotated across adjacent teams and responsibilities."],
            "is_current": False,
        },
    ]


def _career_history_weak(ctx: JDContext, yoe: float, rng: random.Random) -> list:
    # job-hopper shape: many short stints, out-of-domain descriptions.
    n_roles = 4
    per_role_years = max(0.3, round(yoe / n_roles, 1))
    out_terms = ctx.out_of_domain_terms or ["general operations", "administrative support"]
    roles = []
    for i in range(n_roles):
        term = out_terms[i % len(out_terms)]
        roles.append({
            "company": _UNRELATED_COMPANY_POOL[i % len(_UNRELATED_COMPANY_POOL)],
            "title": "Associate",
            "years": per_role_years,
            "bullets": [f"Handled {term}-related tasks and daily coordination."],
            "is_current": i == 0,
        })
    return roles


CAREER_BUILDERS = {
    "strong": _career_history_strong,
    "borderline": _career_history_borderline,
    "weak": _career_history_weak,
}


def _skills_for_tier(ctx: JDContext, tier: str, rng: random.Random) -> list:
    domain_terms = ctx.in_domain_terms or ["general skills"]
    if tier == "strong":
        return domain_terms[:8]
    if tier == "borderline":
        return domain_terms[:3] + ["general project coordination"]
    return ["Microsoft Excel", "Email correspondence", "Scheduling"]


def _education_for_tier(tier: str, rng: random.Random) -> dict:
    if tier == "strong":
        school = rng.choice(_TIER_1_SCHOOLS)
        degree, field_ = "B.Tech", "Computer Science"
    elif tier == "borderline":
        school = "State University"
        degree, field_ = "B.Sc", "General Sciences"
    else:
        school = rng.choice(_TIER_3_SCHOOLS)
        degree, field_ = "Diploma", "Commerce"
    start = rng.randint(2011, 2016)
    return {"institution": school, "degree": degree, "field_of_study": field_,
            "start_year": start, "end_year": start + 4}


def _notice_phrase(tier: str) -> str:
    return {"strong": "Notice period: 15 days",
            "borderline": "Notice period: 60 days",
            "weak": "Notice period: 6 months"}[tier]


def _projects_for_tier(ctx: JDContext, tier: str, rng: random.Random) -> "str | None":
    if tier == "weak":
        return None
    phrases = _signal_phrases(ctx, 2, rng)
    if not phrases:
        return None
    lines = [f"- {_bullet_from_signal(p, rng)} (personal/side project)" for p in phrases]
    return "PROJECTS\n" + "\n".join(lines)


# --------------------------------------------------------------------------- #
# resume text assembly
# --------------------------------------------------------------------------- #
def build_resume(ctx: JDContext, tier: str) -> "ResumeDoc":
    seed_key = f"{ctx.slug}::{tier}"
    rng = random.Random(seed_key)

    yoe = _years_for_tier(ctx, tier, rng)
    location = _location_for_tier(ctx, tier, rng)
    roles = CAREER_BUILDERS[tier](ctx, yoe, rng)
    skills = _skills_for_tier(ctx, tier, rng)
    education = _education_for_tier(tier, rng)
    projects_block = _projects_for_tier(ctx, tier, rng)

    headline = (f"{ctx.title.split(',')[0]}, {yoe:g}+ years" if tier != "weak"
                else "Experienced Professional")
    summary = {
        "strong": f"Results-driven professional with {yoe:g} years of hands-on "
                  f"experience directly aligned to {ctx.domain or ctx.title}. "
                  "Proven track record shipping measurable outcomes.",
        "borderline": f"Generalist professional with {yoe:g} years across "
                      "adjacent roles, with partial exposure to this domain.",
        "weak": f"Professional with {yoe:g} years of experience in unrelated "
                "operational and administrative roles.",
    }[tier]

    return ResumeDoc(
        jd_slug=ctx.slug, tier=tier, headline=headline, summary=summary,
        location=location, years_of_experience=yoe, roles=roles, skills=skills,
        education=education, projects_block=projects_block,
        notice_phrase=_notice_phrase(tier),
    )


_PDF_ASCII_FOLD = {
    "—": "-", "–": "-", "‘": "'", "’": "'",
    "“": '"', "”": '"', "…": "...", " ": " ",
}


def _pdf_safe(text: str) -> str:
    """fpdf2's core Helvetica font is latin-1 only; fold common Unicode
    punctuation (em/en-dash, smart quotes, ellipsis) that JD titles carry,
    then drop anything else outside latin-1 rather than raising."""
    for uni, ascii_ in _PDF_ASCII_FOLD.items():
        text = text.replace(uni, ascii_)
    return text.encode("latin-1", errors="replace").decode("latin-1")


@dataclass
class ResumeDoc:
    jd_slug: str
    tier: str
    headline: str
    summary: str
    location: str
    years_of_experience: float
    roles: list
    skills: list
    education: dict
    projects_block: "str | None"
    notice_phrase: str

    def to_text(self) -> str:
        lines = [
            PLACEHOLDER_NAME,
            f"{PLACEHOLDER_EMAIL} | {PLACEHOLDER_PHONE} | {self.location}, India",
            "",
            "SUMMARY",
            self.summary,
            "",
            "EXPERIENCE",
        ]
        for r in self.roles:
            end = "Present" if r["is_current"] else f"+{r['years']:g}y"
            lines.append(f"{r['title']}, {r['company']} ({r['years']:g} years, {end})")
            for b in r["bullets"]:
                lines.append(f"- {b}")
            lines.append("")
        if self.projects_block:
            lines.append(self.projects_block)
            lines.append("")
        lines.append("EDUCATION")
        e = self.education
        lines.append(f"{e['institution']}, {e['degree']} {e['field_of_study']}, "
                     f"{e['start_year']}-{e['end_year']}")
        lines.append("")
        lines.append("SKILLS")
        lines.append(", ".join(self.skills))
        lines.append("")
        lines.append(self.notice_phrase)
        return "\n".join(lines).strip() + "\n"

    def to_docx_bytes_path(self, path: str) -> None:
        import docx
        doc = docx.Document()
        doc.add_heading(PLACEHOLDER_NAME, level=1)
        doc.add_paragraph(f"{PLACEHOLDER_EMAIL} | {PLACEHOLDER_PHONE} | "
                          f"{self.location}, India")
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(self.summary)
        doc.add_heading("Experience", level=2)
        for r in self.roles:
            end = "Present" if r["is_current"] else f"+{r['years']:g}y"
            doc.add_paragraph(f"{r['title']}, {r['company']} "
                              f"({r['years']:g} years, {end})", style="Heading 3")
            for b in r["bullets"]:
                doc.add_paragraph(b, style="List Bullet")
        if self.projects_block:
            doc.add_heading("Projects", level=2)
            for line in self.projects_block.splitlines()[1:]:
                doc.add_paragraph(line.lstrip("- ").strip(), style="List Bullet")
        doc.add_heading("Education", level=2)
        e = self.education
        doc.add_paragraph(f"{e['institution']}, {e['degree']} {e['field_of_study']}, "
                          f"{e['start_year']}-{e['end_year']}")
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(self.skills))
        doc.add_paragraph(self.notice_phrase)
        doc.save(path)

    def to_pdf_path(self, path: str) -> bool:
        """Write a PDF via fpdf2. Returns False (no exception) if fpdf2 is
        not installed, so callers can skip gracefully."""
        try:
            from fpdf import FPDF, XPos, YPos
        except ImportError:
            return False
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, _pdf_safe(PLACEHOLDER_NAME), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, _pdf_safe(f"{PLACEHOLDER_EMAIL} | {PLACEHOLDER_PHONE} | "
                                  f"{self.location}, India"),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        def heading(text: str) -> None:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, _pdf_safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 10)

        def body(text: str) -> None:
            pdf.multi_cell(0, 6, _pdf_safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        heading("Summary")
        body(self.summary)
        pdf.ln(1)
        heading("Experience")
        for r in self.roles:
            end = "Present" if r["is_current"] else f"+{r['years']:g}y"
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 6, _pdf_safe(f"{r['title']}, {r['company']} "
                                           f"({r['years']:g} years, {end})"),
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 10)
            for b in r["bullets"]:
                body(f"- {b}")
        if self.projects_block:
            pdf.ln(1)
            heading("Projects")
            for line in self.projects_block.splitlines()[1:]:
                body(line)
        pdf.ln(1)
        heading("Education")
        e = self.education
        body(f"{e['institution']}, {e['degree']} {e['field_of_study']}, "
             f"{e['start_year']}-{e['end_year']}")
        pdf.ln(1)
        heading("Skills")
        body(", ".join(self.skills))
        pdf.ln(1)
        body(self.notice_phrase)
        pdf.output(path)
        return True


# --------------------------------------------------------------------------- #
# top-level generation + INDEX.md
# --------------------------------------------------------------------------- #
def _remove_stale_flat_files(base_folder: str) -> None:
    """Pre-layout-dimension runs wrote resume.(txt|docx|pdf) directly into
    `<jd_slug>_<tier>/`. Remove those flat files so re-running this script
    after the layout-matrix change doesn't leave stale duplicates sitting
    next to the new `<layout>/` subfolders. Idempotent / safe to re-run."""
    for name in ("resume.txt", "resume.docx", "resume.pdf"):
        stale = os.path.join(base_folder, name)
        if os.path.isfile(stale):
            os.remove(stale)


def _write_layout_variant(doc: "ResumeDoc", layout: str, folder: str) -> dict:
    """Write one (layout) variant of `doc` into `folder`. Returns a dict
    describing what was written, for the layout-matrix rows in INDEX.md."""
    os.makedirs(folder, exist_ok=True)
    result = {"layout": layout, "folder": folder, "txt_written": False,
              "docx_written": False, "pdf_written": False,
              "ocr_required": layout in resume_layouts.OCR_REQUIRED_LAYOUTS}

    if layout == "single_column":
        txt_path = os.path.join(folder, "resume.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(doc.to_text())
        result["txt_written"] = True

        docx_path = os.path.join(folder, "resume.docx")
        doc.to_docx_bytes_path(docx_path)
        result["docx_written"] = True

        pdf_path = os.path.join(folder, "resume.pdf")
        result["pdf_written"] = doc.to_pdf_path(pdf_path)
        return result

    if layout in resume_layouts.DOCX_RENDERERS:
        txt_path = os.path.join(folder, "resume.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(doc.to_text())
        result["txt_written"] = True

        docx_path = os.path.join(folder, "resume.docx")
        resume_layouts.DOCX_RENDERERS[layout](doc, docx_path)
        result["docx_written"] = True

        pdf_path = os.path.join(folder, "resume.pdf")
        result["pdf_written"] = resume_layouts.PDF_RENDERERS[layout](doc, pdf_path)
        return result

    if layout == "image_only":
        # Derived from the single_column PDF's own output folder, which is
        # always generated first (LAYOUTS order below guarantees this).
        src_pdf = os.path.join(os.path.dirname(folder), "single_column", "resume.pdf")
        pdf_path = os.path.join(folder, "resume.pdf")
        ok = resume_layouts.rasterize_pdf_to_image_only(src_pdf, pdf_path)
        result["pdf_written"] = ok
        if ok:
            result["no_text_layer"] = resume_layouts.pdf_has_no_text(pdf_path)
        return result

    raise ValueError(f"unknown layout: {layout}")


def generate_all(slugs=None, out_dir: str = OUT_DIR) -> list:
    """Generate strong/borderline/weak resumes for each slug, each rendered
    in every layout in `resume_layouts.LAYOUTS`. Returns a list of dicts
    describing what was written (for INDEX.md + CLI summary)."""
    slugs = slugs or TARGET_SLUGS
    os.makedirs(out_dir, exist_ok=True)
    rows = []
    pdf_available = None
    for slug in slugs:
        ctx = load_jd_context(slug)
        for tier in TIERS:
            doc = build_resume(ctx, tier)
            base_folder = os.path.join(out_dir, f"{slug}_{tier}")
            os.makedirs(base_folder, exist_ok=True)
            _remove_stale_flat_files(base_folder)

            for layout in resume_layouts.LAYOUTS:
                layout_folder = os.path.join(base_folder, layout)
                variant = _write_layout_variant(doc, layout, layout_folder)
                if pdf_available is None and layout != "image_only":
                    pdf_available = variant["pdf_written"]

                rows.append({
                    "jd_slug": slug, "tier": tier, "layout": layout,
                    "folder": layout_folder,
                    "years_of_experience": doc.years_of_experience,
                    "location": doc.location,
                    "txt_written": variant["txt_written"],
                    "docx_written": variant["docx_written"],
                    "pdf_written": variant["pdf_written"],
                    "ocr_required": variant["ocr_required"],
                    "no_text_layer": variant.get("no_text_layer"),
                })
    if pdf_available is False:
        print("[gen_resumes] WARNING: fpdf2 not installed — PDF export skipped "
              "for all resumes (txt + docx still written). "
              "`pip install -r requirements-db.txt` to enable PDF export.")
    _write_index(rows, out_dir)
    return rows


def _write_index(rows: list, out_dir: str) -> None:
    lines = [
        "# Synthetic resumes — INDEX",
        "",
        "Generated by `tools/gen_resumes.py` (Phase 2, Part C3). Placeholder PII "
        f"only ({PLACEHOLDER_NAME} / {PLACEHOLDER_EMAIL} / {PLACEHOLDER_PHONE}); "
        "deterministic via `random.seed(f\"{jd_slug}::{tier}\")`. Each resume is "
        "seeded off its gold JD's `jd_meta.yaml` must_haves/signals so fit tiers "
        "(strong/borderline/weak) are realistic against "
        "`redrob_ranker.fit.score_candidate`.",
        "",
        "## Layout matrix",
        "",
        "Each (jd_slug, tier) is rendered in every layout below — same "
        "content/PII/seed, only the rendering (and file formats) vary. See "
        "`tools/resume_layouts.py`.",
        "",
        "| layout | description | formats | requires OCR |",
        "|---|---|---|---|",
        "| single_column | clean single-column baseline | txt, docx, pdf | no |",
        "| two_column | sidebar (contact/skills/education) + main "
        "(summary/experience) | txt, docx, pdf | no |",
        "| table_heavy | experience + skills as bordered tables/grids | "
        "txt, docx, pdf | no |",
        "| ats_plain | minimal, no styling, plain headings (ATS-safe baseline) "
        "| txt, docx, pdf | no |",
        "| image_only | single_column PDF rasterized into a NEW PDF with **no "
        "text layer** | pdf only | **yes** |",
        "",
        "## Files",
        "",
        "| jd_slug | tier | layout | years_of_experience | location | folder | "
        "txt | docx | pdf | ocr_required |",
        "|---|---|---|---|---|---|---|---|---|---|",
    ]
    for r in rows:
        rel = os.path.relpath(r["folder"], out_dir).replace(os.sep, "/")
        txt_mark = "yes" if r["txt_written"] else "n/a"
        docx_mark = "yes" if r["docx_written"] else "n/a"
        if r["pdf_written"]:
            pdf_mark = "yes"
            if r.get("no_text_layer") is True:
                pdf_mark = "yes (no text layer)"
            elif r.get("no_text_layer") is False:
                pdf_mark = "yes (WARNING: text layer detected)"
        else:
            pdf_mark = "skipped (fpdf2/pymupdf absent)"
        ocr_mark = "**yes**" if r["ocr_required"] else "no"
        lines.append(f"| {r['jd_slug']} | {r['tier']} | {r['layout']} | "
                     f"{r['years_of_experience']:g} | {r['location']} | "
                     f"{rel}/ | {txt_mark} | {docx_mark} | {pdf_mark} | {ocr_mark} |")
    with open(os.path.join(out_dir, "INDEX.md"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


# --------------------------------------------------------------------------- #
# CLI: python -m tools.gen_resumes
# --------------------------------------------------------------------------- #
def _main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Generate synthetic resumes for the gold JDs.")
    ap.add_argument("--slugs", nargs="*", default=None,
                    help="jd_slugs to generate for (default: 4-JD spread, 12 resumes)")
    ap.add_argument("--out", default=OUT_DIR)
    args = ap.parse_args(argv)

    rows = generate_all(slugs=args.slugs, out_dir=args.out)
    n_resumes = len({(r["jd_slug"], r["tier"]) for r in rows})
    n_files = sum(int(r["txt_written"]) + int(r["docx_written"]) + int(r["pdf_written"])
                  for r in rows)
    print(f"[gen_resumes] wrote {len(rows)} layout variants "
          f"({n_resumes} resumes x {len(resume_layouts.LAYOUTS)} layouts), "
          f"{n_files} files total, to {args.out}")
    for r in rows:
        pdf_note = "ok"
        if not r["pdf_written"]:
            pdf_note = "skipped"
        elif r.get("no_text_layer") is True:
            pdf_note = "ok (no text layer)"
        print(f"  {r['jd_slug']}_{r['tier']}/{r['layout']}: "
              f"yoe={r['years_of_experience']:g} loc={r['location']} pdf={pdf_note}")
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(_main())
