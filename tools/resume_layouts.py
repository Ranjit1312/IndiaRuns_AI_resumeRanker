"""resume_layouts.py — structural-variety renderers for synthetic resumes.

`tools/gen_resumes.py` builds one `ResumeDoc` per (jd_slug, tier). This module
takes that SAME content/PII/seed and re-renders it into several common
real-world resume *structures*, so the JD-parsing / OCR pipeline gets tested
against more than one clean single-column shape.

Layouts:
  - single_column — the original layout, unchanged (see gen_resumes.ResumeDoc).
  - two_column    — sidebar (contact/skills/education) + main (summary/experience).
  - table_heavy   — experience + skills rendered as bordered tables/grids.
  - ats_plain     — minimal plain-text-like styling (the ATS-friendly baseline).
  - image_only    — the single_column PDF's pages rasterized into a NEW PDF
                    with no text layer at all (requires OCR to read). Built
                    from an already-rendered single_column PDF via pymupdf;
                    it has no independent DOCX/txt form.

Every function here is a pure renderer: given a `ResumeDoc` (defined in
tools.gen_resumes) and an output path, it writes that one file. No content
generation, no randomness, no PII — all of that already happened upstream in
`build_resume`. This keeps the layout dimension orthogonal to the content/tier
dimension, and keeps re-runs byte-identical (same seed in -> same bytes out).
"""
from __future__ import annotations

LAYOUTS = ["single_column", "two_column", "table_heavy", "ats_plain", "image_only"]

# Layouts that have a meaningful plain-text rendering. image_only is a raster
# PDF with no text layer by design, so a .txt sibling would be misleading.
TEXT_LAYOUTS = ["single_column", "two_column", "table_heavy", "ats_plain"]

# Layouts that exist only as a PDF derived from another layout's PDF.
DERIVED_PDF_ONLY_LAYOUTS = ["image_only"]

OCR_REQUIRED_LAYOUTS = {"image_only"}


# --------------------------------------------------------------------------- #
# shared helpers
# --------------------------------------------------------------------------- #
def _role_line(r: dict) -> str:
    end = "Present" if r["is_current"] else f"+{r['years']:g}y"
    return f"{r['title']}, {r['company']} ({r['years']:g} years, {end})"


def _edu_line(e: dict) -> str:
    return (f"{e['institution']}, {e['degree']} {e['field_of_study']}, "
            f"{e['start_year']}-{e['end_year']}")


def _contact_line(doc) -> str:
    from tools.gen_resumes import PLACEHOLDER_EMAIL, PLACEHOLDER_PHONE
    return f"{PLACEHOLDER_EMAIL} | {PLACEHOLDER_PHONE} | {doc.location}, India"


# --------------------------------------------------------------------------- #
# two_column
# --------------------------------------------------------------------------- #
def two_column_to_docx(doc, path: str) -> None:
    """Sidebar (contact/skills/education) + main column (summary/experience),
    implemented as a 2-column, 1-row DOCX table (no visible borders)."""
    import docx
    from docx.shared import Inches
    from tools.gen_resumes import PLACEHOLDER_NAME

    d = docx.Document()
    d.add_heading(PLACEHOLDER_NAME, level=1)
    d.add_paragraph(_contact_line(doc))

    table = d.add_table(rows=1, cols=2)
    table.autofit = False
    sidebar_cell, main_cell = table.rows[0].cells
    sidebar_cell.width = Inches(2.0)
    main_cell.width = Inches(4.5)

    # sidebar: skills, education
    sidebar_cell.paragraphs[0].add_run("Skills").bold = True
    for s in doc.skills:
        sidebar_cell.add_paragraph(s, style="List Bullet")
    sidebar_cell.add_paragraph("").add_run("Education").bold = True
    e = doc.education
    sidebar_cell.add_paragraph(_edu_line(e))
    sidebar_cell.add_paragraph(doc.notice_phrase)

    # main: summary, experience, projects
    main_cell.paragraphs[0].add_run("Summary").bold = True
    main_cell.add_paragraph(doc.summary)
    p = main_cell.add_paragraph()
    p.add_run("Experience").bold = True
    for r in doc.roles:
        rp = main_cell.add_paragraph()
        rp.add_run(_role_line(r)).bold = True
        for b in r["bullets"]:
            main_cell.add_paragraph(b, style="List Bullet")
    if doc.projects_block:
        pp = main_cell.add_paragraph()
        pp.add_run("Projects").bold = True
        for line in doc.projects_block.splitlines()[1:]:
            main_cell.add_paragraph(line.lstrip("- ").strip(), style="List Bullet")

    d.save(path)


def two_column_to_pdf(doc, path: str) -> bool:
    """Two text columns via multi_cell with explicit x offsets: a narrow left
    sidebar (contact/skills/education) and a wider right main column
    (summary/experience/projects)."""
    try:
        from fpdf import FPDF, XPos, YPos
    except ImportError:
        return False
    from tools.gen_resumes import PLACEHOLDER_NAME, _pdf_safe

    LEFT_X, LEFT_W = 10, 60
    RIGHT_X, RIGHT_W = 75, 125
    TOP_Y = 30

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_xy(10, 10)
    pdf.cell(0, 10, _pdf_safe(PLACEHOLDER_NAME), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_xy(10, 20)
    pdf.multi_cell(0, 6, _pdf_safe(_contact_line(doc)))

    def col_heading(text: str, x: float, y: float) -> float:
        pdf.set_xy(x, y)
        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(LEFT_W if x == LEFT_X else RIGHT_W, 6, _pdf_safe(text))
        return pdf.get_y()

    def col_body(text: str, x: float, y: float, w: float) -> float:
        pdf.set_xy(x, y)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(w, 5, _pdf_safe(text))
        return pdf.get_y()

    # left sidebar
    y = col_heading("Skills", LEFT_X, TOP_Y)
    y = col_body(", ".join(doc.skills), LEFT_X, y, LEFT_W)
    y = col_heading("Education", LEFT_X, y + 3)
    y = col_body(_edu_line(doc.education), LEFT_X, y, LEFT_W)
    col_body(doc.notice_phrase, LEFT_X, y + 3, LEFT_W)

    # right main column
    y = col_heading("Summary", RIGHT_X, TOP_Y)
    y = col_body(doc.summary, RIGHT_X, y, RIGHT_W)
    y = col_heading("Experience", RIGHT_X, y + 3)
    for r in doc.roles:
        pdf.set_xy(RIGHT_X, y)
        pdf.set_font("Helvetica", "B", 9)
        pdf.multi_cell(RIGHT_W, 5, _pdf_safe(_role_line(r)))
        y = pdf.get_y()
        for b in r["bullets"]:
            y = col_body(f"- {b}", RIGHT_X, y, RIGHT_W)
    if doc.projects_block:
        y = col_heading("Projects", RIGHT_X, y + 3)
        for line in doc.projects_block.splitlines()[1:]:
            y = col_body(line, RIGHT_X, y, RIGHT_W)

    pdf.output(path)
    return True


# --------------------------------------------------------------------------- #
# table_heavy
# --------------------------------------------------------------------------- #
def table_heavy_to_docx(doc, path: str) -> None:
    """Experience and skills rendered as bordered DOCX tables."""
    import docx
    from tools.gen_resumes import PLACEHOLDER_NAME

    d = docx.Document()
    d.add_heading(PLACEHOLDER_NAME, level=1)
    d.add_paragraph(_contact_line(doc))
    d.add_heading("Summary", level=2)
    d.add_paragraph(doc.summary)

    d.add_heading("Experience", level=2)
    exp_table = d.add_table(rows=1, cols=3)
    exp_table.style = "Table Grid"
    hdr = exp_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Role / Company", "Duration", "Highlights"
    for r in doc.roles:
        row = exp_table.add_row().cells
        row[0].text = f"{r['title']}\n{r['company']}"
        end = "Present" if r["is_current"] else f"+{r['years']:g}y"
        row[1].text = f"{r['years']:g} years ({end})"
        row[2].text = "\n".join(f"- {b}" for b in r["bullets"])

    if doc.projects_block:
        d.add_heading("Projects", level=2)
        for line in doc.projects_block.splitlines()[1:]:
            d.add_paragraph(line.lstrip("- ").strip(), style="List Bullet")

    d.add_heading("Education", level=2)
    d.add_paragraph(_edu_line(doc.education))

    d.add_heading("Skills", level=2)
    skills_table = d.add_table(rows=0, cols=4)
    skills_table.style = "Table Grid"
    row_cells = None
    for i, s in enumerate(doc.skills):
        if i % 4 == 0:
            row_cells = skills_table.add_row().cells
        row_cells[i % 4].text = s

    d.add_paragraph(doc.notice_phrase)
    d.save(path)


def table_heavy_to_pdf(doc, path: str) -> bool:
    """Experience + skills rendered as bordered cell grids."""
    try:
        from fpdf import FPDF, XPos, YPos
    except ImportError:
        return False
    from tools.gen_resumes import PLACEHOLDER_NAME, _pdf_safe

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _pdf_safe(PLACEHOLDER_NAME), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 8, _pdf_safe(_contact_line(doc)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    def heading(text: str) -> None:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, _pdf_safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)

    heading("Summary")
    pdf.multi_cell(0, 6, _pdf_safe(doc.summary), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)

    heading("Experience")
    col_w = [55, 30, 105]
    row_h = 6
    pdf.set_font("Helvetica", "B", 9)
    for w, label in zip(col_w, ["Role / Company", "Duration", "Highlights"]):
        pdf.cell(w, row_h, label, border=1)
    pdf.ln(row_h)
    pdf.set_font("Helvetica", "", 9)
    for r in doc.roles:
        end = "Present" if r["is_current"] else f"+{r['years']:g}y"
        role_txt = _pdf_safe(f"{r['title']}, {r['company']}")[:38]
        dur_txt = _pdf_safe(f"{r['years']:g}y ({end})")
        highlights = _pdf_safe("; ".join(r["bullets"]))[:80]
        x0, y0 = pdf.get_x(), pdf.get_y()
        pdf.multi_cell(col_w[0], row_h, role_txt, border=1)
        y1 = pdf.get_y()
        pdf.set_xy(x0 + col_w[0], y0)
        pdf.multi_cell(col_w[1], row_h, dur_txt, border=1)
        pdf.set_xy(x0 + col_w[0] + col_w[1], y0)
        pdf.multi_cell(col_w[2], row_h, highlights, border=1)
        y2 = pdf.get_y()
        pdf.set_xy(x0, max(y1, y2))
    pdf.ln(2)

    if doc.projects_block:
        heading("Projects")
        for line in doc.projects_block.splitlines()[1:]:
            pdf.multi_cell(0, 6, _pdf_safe(line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    heading("Education")
    pdf.multi_cell(0, 6, _pdf_safe(_edu_line(doc.education)),
                    new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)

    heading("Skills")
    pdf.set_font("Helvetica", "", 9)
    cell_w = 45
    for i, s in enumerate(doc.skills):
        if i and i % 4 == 0:
            pdf.ln(row_h)
        pdf.cell(cell_w, row_h, _pdf_safe(s)[:22], border=1)
    pdf.ln(row_h + 2)

    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, _pdf_safe(doc.notice_phrase), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(path)
    return True


# --------------------------------------------------------------------------- #
# ats_plain — minimal, no styling, plain headings (ATS-friendly baseline)
# --------------------------------------------------------------------------- #
def ats_plain_to_docx(doc, path: str) -> None:
    """No heading styles, no bullet lists, no bold — just plain paragraphs
    with all-caps section labels, the way an ATS-safe template looks."""
    import docx
    from tools.gen_resumes import PLACEHOLDER_NAME

    d = docx.Document()
    d.add_paragraph(PLACEHOLDER_NAME)
    d.add_paragraph(_contact_line(doc))
    d.add_paragraph("")
    d.add_paragraph("SUMMARY")
    d.add_paragraph(doc.summary)
    d.add_paragraph("")
    d.add_paragraph("EXPERIENCE")
    for r in doc.roles:
        d.add_paragraph(_role_line(r))
        for b in r["bullets"]:
            d.add_paragraph(f"- {b}")
    if doc.projects_block:
        d.add_paragraph("")
        for line in doc.projects_block.splitlines():
            d.add_paragraph(line)
    d.add_paragraph("")
    d.add_paragraph("EDUCATION")
    d.add_paragraph(_edu_line(doc.education))
    d.add_paragraph("")
    d.add_paragraph("SKILLS")
    d.add_paragraph(", ".join(doc.skills))
    d.add_paragraph("")
    d.add_paragraph(doc.notice_phrase)
    d.save(path)


def ats_plain_to_pdf(doc, path: str) -> bool:
    """Single font, single size, no bold headings — plain text laid out top
    to bottom, exactly what an ATS parser expects to see."""
    try:
        from fpdf import FPDF, XPos, YPos
    except ImportError:
        return False
    from tools.gen_resumes import PLACEHOLDER_NAME, _pdf_safe

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 11)

    def line(text: str) -> None:
        pdf.multi_cell(0, 6, _pdf_safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    line(PLACEHOLDER_NAME)
    line(_contact_line(doc))
    line("")
    line("SUMMARY")
    line(doc.summary)
    line("")
    line("EXPERIENCE")
    for r in doc.roles:
        line(_role_line(r))
        for b in r["bullets"]:
            line(f"- {b}")
    if doc.projects_block:
        line("")
        for l in doc.projects_block.splitlines():
            line(l)
    line("")
    line("EDUCATION")
    line(_edu_line(doc.education))
    line("")
    line("SKILLS")
    line(", ".join(doc.skills))
    line("")
    line(doc.notice_phrase)

    pdf.output(path)
    return True


# --------------------------------------------------------------------------- #
# image_only — rasterize an existing text PDF into a new, text-free PDF
# --------------------------------------------------------------------------- #
def rasterize_pdf_to_image_only(src_pdf_path: str, out_pdf_path: str,
                                 dpi: int = 200) -> bool:
    """Render every page of `src_pdf_path` to a raster image and rebuild a NEW
    PDF containing ONLY those page images (no selectable text layer at all).
    This is the scanned-resume simulation that forces the OCR path.

    Returns False (no exception) if pymupdf is unavailable, so callers can
    skip gracefully. Requires `src_pdf_path` to already exist.
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        return False
    if not __import__("os").path.exists(src_pdf_path):
        return False

    src = fitz.open(src_pdf_path)
    out = fitz.open()
    for page in src:
        pix = page.get_pixmap(dpi=dpi)
        img_pdf_bytes = pix.pil_tobytes(format="PNG") if hasattr(pix, "pil_tobytes") else pix.tobytes("png")
        rect = fitz.Rect(0, 0, page.rect.width, page.rect.height)
        new_page = out.new_page(width=rect.width, height=rect.height)
        new_page.insert_image(rect, stream=img_pdf_bytes)
    out.save(out_pdf_path)
    out.close()
    src.close()
    return True


def pdf_has_no_text(pdf_path: str) -> bool:
    """True if every page of `pdf_path` has an empty extractable text layer
    (i.e. it truly requires OCR). Used both by the generator's own
    self-check and by tests."""
    import fitz  # pymupdf
    d = fitz.open(pdf_path)
    try:
        return all(page.get_text().strip() == "" for page in d)
    finally:
        d.close()


# --------------------------------------------------------------------------- #
# dispatch tables consumed by gen_resumes.generate_all
# --------------------------------------------------------------------------- #
DOCX_RENDERERS = {
    "two_column": two_column_to_docx,
    "table_heavy": table_heavy_to_docx,
    "ats_plain": ats_plain_to_docx,
}

PDF_RENDERERS = {
    "two_column": two_column_to_pdf,
    "table_heavy": table_heavy_to_pdf,
    "ats_plain": ats_plain_to_pdf,
}
